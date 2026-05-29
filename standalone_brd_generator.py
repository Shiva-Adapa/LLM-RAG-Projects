import streamlit as st
import os
import re
from io import StringIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import AzureOpenAI

# ========== CONFIG ========== #
# Azure OpenAI config
AZURE_OPENAI_ENDPOINT = "**************************************************************"
AZURE_OPENAI_KEY = "**************************************************************"
AZURE_DEPLOYMENT_NAME = "gpt-5-chat"  

client = AzureOpenAI(
    api_key=AZURE_OPENAI_KEY,
    api_version="2024-05-01-preview",
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)

st.set_page_config(page_title="BRD Generator", layout="wide")
st.title("📄 Business Requirement Document (BRD) Generator")

uploaded_transcript = st.file_uploader("Upload meeting transcript (TXT file)", type=["txt"])
uploaded_template = st.file_uploader("Optional: Upload BRD template (TXT/MD/DOCX)", type=["txt", "md", "docx"])

# ========== READ FILES ========== #
transcript_text = None
if uploaded_transcript is not None:
    stringio = StringIO(uploaded_transcript.getvalue().decode("utf-8"))
    transcript_text = stringio.read()

template_text = None
if uploaded_template is not None:
    if uploaded_template.name.endswith((".txt", ".md")):
        stringio = StringIO(uploaded_template.getvalue().decode("utf-8"))
        template_text = stringio.read()
    elif uploaded_template.name.endswith(".docx"):
        doc = Document(uploaded_template)
        template_text = "\n".join([p.text for p in doc.paragraphs])

# ========== CLEANING FUNCTIONS ========== #
def clean_brd_text(raw_text: str) -> str:
    """Remove chatty lines and only keep formal BRD content."""
    # Keep only after first BRD heading
    match = re.search(r"(#\s*\*\*?Business Requirements Document.*)", raw_text, re.IGNORECASE | re.DOTALL)
    if match:
        raw_text = match.group(1)

    # Remove filler sentences
    patterns_to_remove = [
        r"(?i)would you like.*", 
        r"(?i)do you want me.*",
        r"(?i)got it.*",
        r"(?i)i['’]ll.*",
        r"(?i)if you’d like.*"
    ]
    for pattern in patterns_to_remove:
        raw_text = re.sub(pattern, "", raw_text)

    # Remove extra blank lines
    raw_text = re.sub(r"\n\s*\n\s*\n+", "\n\n", raw_text).strip()
    return raw_text

def add_paragraph_with_bold(doc, line):
    """Convert markdown bold **text** into Word bold formatting."""
    bold_matches = re.findall(r"\*\*(.*?)\*\*", line)
    if bold_matches:
        para = doc.add_paragraph()
        cursor = 0
        for match in bold_matches:
            start = line.find(f"**{match}**", cursor)
            if start > cursor:
                para.add_run(line[cursor:start])
            para.add_run(match).bold = True
            cursor = start + len(f"**{match}**")
        if cursor < len(line):
            para.add_run(line[cursor:])
    else:
        doc.add_paragraph(line)

def extract_dates_from_text(text):
    """Extract dates, timelines, TATs mentioned in transcript."""
    # Simple regex for dates like Sept 1, October 31, 2025, etc.
    dates = re.findall(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2}(?:,?\s*\d{4})?)', text, re.IGNORECASE)
    return list(set(dates))  # unique dates

# ========== SAVE BRD FUNCTION ========== #
def save_brd_to_word(brd_text: str, transcript_text: str):
    cleaned_text = clean_brd_text(brd_text)
    extracted_dates = extract_dates_from_text(transcript_text)

    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    brd_folder = os.path.join(desktop, "brd")
    os.makedirs(brd_folder, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = os.path.join(brd_folder, f"BRD_{timestamp}.docx")

    doc = Document()
    doc.add_heading("Business Requirement Document", level=0)

    # ========== CONTEXT TABLE ========== #
    context_table = doc.add_table(rows=1, cols=2)
    context_table.style = 'Table Grid'
    hdr_cells = context_table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Description"

    context_data = [
        ("Project Name", "ProLift Tire Sale Customer Targeting Report"),
        ("Project Sponsor", "[Name / Department]"),
        ("Business Owner", "[Name / Department]"),
        ("Target System", "Sales & Invoicing System"),
        ("Data Source", "Customer purchase history for past 5 years"),
        ("Report Users", "Sales Specialists, Marketing Team"),
        ("Date", datetime.now().strftime("%B %d, %Y"))
    ]
    for item, desc in context_data:
        row_cells = context_table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc

    # ========== PROCESS BRD LINES ========== #
    skip_sections = ["Out of Scope", "Non-Functional Requirements"]
    current_skip = False
    for line in cleaned_text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Skip sections if mentioned in skip_sections and not present in text
        for section in skip_sections:
            if line.lower().startswith(section.lower()):
                current_skip = section not in transcript_text
        if current_skip:
            continue

        if line.startswith("##") or line.startswith("#"):
            doc.add_heading(line.strip("# ").strip(), level=1)
        elif "|" in line and "---" not in line:
            # Table
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if not hasattr(doc, "_last_table") or len(cells) != len(doc._last_table.rows[0].cells):
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, c in enumerate(cells):
                    hdr_cells[i].text = c
                doc._last_table = table
            else:
                row_cells = doc._last_table.add_row().cells
                for i, c in enumerate(cells):
                    row_cells[i].text = c
        else:
            add_paragraph_with_bold(doc, line)

    doc.save(filename)
    return filename

# ========== STREAMLIT BUTTON ========== #
if st.button("Generate BRD"):
    if not transcript_text:
        st.error("Please upload a meeting transcript first.")
    else:
        with st.spinner("Generating BRD using Azure OpenAI GPT-5..."):
            if template_text:
                system_prompt = f"""You are a senior business analyst.
Convert the following meeting transcript into a formal BRD strictly following this template:
{template_text}"""
            else:
                system_prompt = """You are a senior business analyst.
Convert the following meeting transcript into a formal BRD.
Use numbered sections (Introduction, Objectives, Scope, Stakeholders, Functional & Non-Functional Requirements, Assumptions, Constraints, Timeline, Sign-off).
Exclude any sections not mentioned in the transcript (like Out of Scope or Non-Functional Requirements).
Extract any dates, deadlines, timelines, TAT mentioned in the text.
Do NOT include suggestions, chatty phrases, or filler text.
Use **bold** for headings and important text, and tables for requirements/timeline.
"""

            response = client.chat.completions.create(
                model=AZURE_DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": transcript_text}
                ],
                temperature=0.2,
                max_tokens=2500
            )

            brd_text = response.choices[0].message.content
            file_path = save_brd_to_word(brd_text, transcript_text)

            st.subheader("📑 Generated BRD")
            st.text_area("Generated Document", brd_text, height=600)
            st.download_button(
                label="Download BRD as TXT",
                data=brd_text,
                file_name="Generated_BRD.txt",
                mime="text/plain"
            )
            st.success(f"✅ BRD saved to: {file_path}")
